"""
Enterprise Report Generation Service for ScoutAI
"""
import logging
import os
import io
from datetime import datetime
from typing import Dict, List, Any, Optional
import pandas as pd
import numpy as np

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

import matplotlib.pyplot as plt
import matplotlib.patches as patches
import seaborn as sns
from matplotlib.backends.backend_pdf import PdfPages
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.chart import BarChart, LineChart, Reference

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..database import db_manager
from ..services.llm_service import llm_service
from ..models.enhanced_ml_models import player_model, team_model

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Enterprise-grade report generation service"""
    
    def __init__(self):
        self.report_dir = "reports"
        os.makedirs(self.report_dir, exist_ok=True)
        
        # Set up matplotlib style
        plt.style.use('seaborn-v0_8-darkgrid')
        sns.set_palette("husl")
    
    def generate_player_scouting_report(self, player_id: int, format: str = "pdf") -> str:
        """Generate comprehensive player scouting report"""
        try:
            # Get player data
            player_data = self._get_player_data(player_id)
            if not player_data:
                raise ValueError(f"No data found for player ID: {player_id}")
            
            # Get LLM analysis
            llm_analysis = llm_service.generate_player_analysis(
                player_data['personal'], 
                player_data['performance']
            )
            
            # Get ML predictions
            ml_predictions = player_model.predict_player_potential(player_id)
            
            # Generate report based on format
            if format.lower() == "pdf":
                return self._generate_player_pdf_report(player_data, llm_analysis, ml_predictions)
            elif format.lower() == "excel":
                return self._generate_player_excel_report(player_data, llm_analysis, ml_predictions)
            elif format.lower() == "word":
                return self._generate_player_word_report(player_data, llm_analysis, ml_predictions)
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            logger.error(f"Error generating player report: {e}")
            raise
    
    def generate_team_analysis_report(self, team_name: str, season: int = None, format: str = "pdf") -> str:
        """Generate comprehensive team analysis report"""
        try:
            # Get team data
            team_data = self._get_team_data(team_name, season)
            if not team_data:
                raise ValueError(f"No data found for team: {team_name}")
            
            # Get ML analysis
            ml_analysis = team_model.analyze_team_performance(team_name, season)
            
            # Get LLM insights
            llm_analysis = llm_service.generate_team_analysis(team_data, team_data.get('recent_matches', []))
            
            # Generate report based on format
            if format.lower() == "pdf":
                return self._generate_team_pdf_report(team_data, ml_analysis, llm_analysis)
            elif format.lower() == "excel":
                return self._generate_team_excel_report(team_data, ml_analysis, llm_analysis)
            elif format.lower() == "word":
                return self._generate_team_word_report(team_data, ml_analysis, llm_analysis)
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            logger.error(f"Error generating team report: {e}")
            raise
    
    def generate_league_summary_report(self, season: int, format: str = "pdf") -> str:
        """Generate league-wide summary report"""
        try:
            # Get league data
            league_data = self._get_league_data(season)
            
            # Generate report based on format
            if format.lower() == "pdf":
                return self._generate_league_pdf_report(league_data, season)
            elif format.lower() == "excel":
                return self._generate_league_excel_report(league_data, season)
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            logger.error(f"Error generating league report: {e}")
            raise
    
    def _generate_player_pdf_report(self, player_data: Dict, llm_analysis: str, ml_predictions: Dict) -> str:
        """Generate PDF player scouting report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/player_report_{player_data['personal']['first_name']}_{player_data['personal']['last_name']}_{timestamp}.pdf"
        
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        player_name = f"{player_data['personal']['first_name']} {player_data['personal']['last_name']}"
        story.append(Paragraph(f"PLAYER SCOUTING REPORT: {player_name}", title_style))
        story.append(Spacer(1, 20))
        
        # Player Overview Table
        personal_info = player_data['personal']
        overview_data = [
            ['Full Name', player_name],
            ['Date of Birth', str(personal_info.get('born_date', 'N/A'))],
            ['Height', f"{personal_info.get('height', 'N/A')} cm"],
            ['Weight', f"{personal_info.get('weight', 'N/A')} kg"],
            ['Debut Date', str(personal_info.get('debut_date', 'N/A'))],
        ]
        
        overview_table = Table(overview_data, colWidths=[2*inch, 3*inch])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(Paragraph("Player Overview", styles['Heading2']))
        story.append(overview_table)
        story.append(Spacer(1, 20))
        
        # Performance Statistics
        if player_data['performance']:
            story.append(Paragraph("Performance Analysis", styles['Heading2']))
            
            # Calculate career averages
            perf_df = pd.DataFrame(player_data['performance'])
            career_avg = self._calculate_career_averages(perf_df)
            
            stats_data = [['Statistic', 'Career Average']]
            for stat, value in career_avg.items():
                stats_data.append([stat.replace('_', ' ').title(), f"{value:.2f}"])
            
            stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(stats_table)
            story.append(Spacer(1, 20))
        
        # ML Predictions
        if ml_predictions and 'error' not in ml_predictions:
            story.append(Paragraph("AI-Powered Analysis", styles['Heading2']))
            
            pred_data = [
                ['Current Impact Score', f"{ml_predictions.get('current_impact_score', 0):.2f}"],
                ['Career Stage', ml_predictions.get('career_stage', 'Unknown')],
                ['Injury Risk', ml_predictions.get('injury_risk', 'Unknown')],
            ]
            
            pred_table = Table(pred_data, colWidths=[3*inch, 2*inch])
            pred_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(pred_table)
            story.append(Spacer(1, 20))
        
        # LLM Analysis
        story.append(Paragraph("Expert Scouting Analysis", styles['Heading2']))
        story.append(Paragraph(llm_analysis, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        logger.info(f"Generated player PDF report: {filename}")
        
        return filename
    
    def _generate_player_excel_report(self, player_data: Dict, llm_analysis: str, ml_predictions: Dict) -> str:
        """Generate Excel player scouting report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/player_report_{player_data['personal']['first_name']}_{player_data['personal']['last_name']}_{timestamp}.xlsx"
        
        wb = Workbook()
        
        # Player Overview Sheet
        ws1 = wb.active
        ws1.title = "Player Overview"
        
        # Headers
        headers = ['Attribute', 'Value']
        for col, header in enumerate(headers, 1):
            cell = ws1.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            cell.font = Font(color="FFFFFF", bold=True)
        
        # Player data
        personal = player_data['personal']
        player_info = [
            ['Full Name', f"{personal['first_name']} {personal['last_name']}"],
            ['Date of Birth', str(personal.get('born_date', 'N/A'))],
            ['Height (cm)', personal.get('height', 'N/A')],
            ['Weight (kg)', personal.get('weight', 'N/A')],
            ['Debut Date', str(personal.get('debut_date', 'N/A'))],
        ]
        
        for row, (attr, value) in enumerate(player_info, 2):
            ws1.cell(row=row, column=1, value=attr)
            ws1.cell(row=row, column=2, value=value)
        
        # Performance Data Sheet
        if player_data['performance']:
            ws2 = wb.create_sheet("Performance Data")
            
            perf_df = pd.DataFrame(player_data['performance'])
            
            # Write headers
            for col, header in enumerate(perf_df.columns, 1):
                cell = ws2.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
                cell.font = Font(color="FFFFFF", bold=True)
            
            # Write data
            for row, data_row in enumerate(perf_df.values, 2):
                for col, value in enumerate(data_row, 1):
                    ws2.cell(row=row, column=col, value=value)
        
        # Analysis Sheet
        ws3 = wb.create_sheet("AI Analysis")
        
        # LLM Analysis
        ws3.cell(row=1, column=1, value="Expert Analysis").font = Font(bold=True, size=14)
        ws3.cell(row=2, column=1, value=llm_analysis)
        ws3.cell(row=2, column=1).alignment = Alignment(wrap_text=True)
        
        # ML Predictions
        if ml_predictions and 'error' not in ml_predictions:
            ws3.cell(row=10, column=1, value="ML Predictions").font = Font(bold=True, size=14)
            
            pred_info = [
                ['Current Impact Score', f"{ml_predictions.get('current_impact_score', 0):.2f}"],
                ['Career Stage', ml_predictions.get('career_stage', 'Unknown')],
                ['Injury Risk', ml_predictions.get('injury_risk', 'Unknown')],
            ]
            
            for row, (attr, value) in enumerate(pred_info, 11):
                ws3.cell(row=row, column=1, value=attr)
                ws3.cell(row=row, column=2, value=value)
        
        # Auto-adjust column widths
        for ws in wb.worksheets:
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(filename)
        logger.info(f"Generated player Excel report: {filename}")
        
        return filename
    
    def _generate_player_word_report(self, player_data: Dict, llm_analysis: str, ml_predictions: Dict) -> str:
        """Generate Word player scouting report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/player_report_{player_data['personal']['first_name']}_{player_data['personal']['last_name']}_{timestamp}.docx"
        
        doc = Document()
        
        # Title
        player_name = f"{player_data['personal']['first_name']} {player_data['personal']['last_name']}"
        title = doc.add_heading(f'Player Scouting Report: {player_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Player Overview
        doc.add_heading('Player Overview', level=1)
        
        personal = player_data['personal']
        overview_table = doc.add_table(rows=6, cols=2)
        overview_table.style = 'Table Grid'
        
        overview_data = [
            ['Attribute', 'Value'],
            ['Full Name', player_name],
            ['Date of Birth', str(personal.get('born_date', 'N/A'))],
            ['Height', f"{personal.get('height', 'N/A')} cm"],
            ['Weight', f"{personal.get('weight', 'N/A')} kg"],
            ['Debut Date', str(personal.get('debut_date', 'N/A'))],
        ]
        
        for row_idx, (attr, value) in enumerate(overview_data):
            overview_table.rows[row_idx].cells[0].text = attr
            overview_table.rows[row_idx].cells[1].text = str(value)
            if row_idx == 0:  # Header row
                for cell in overview_table.rows[row_idx].cells:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # Performance Analysis
        if player_data['performance']:
            doc.add_heading('Performance Statistics', level=1)
            
            perf_df = pd.DataFrame(player_data['performance'])
            career_avg = self._calculate_career_averages(perf_df)
            
            stats_table = doc.add_table(rows=len(career_avg) + 1, cols=2)
            stats_table.style = 'Table Grid'
            
            # Headers
            stats_table.rows[0].cells[0].text = 'Statistic'
            stats_table.rows[0].cells[1].text = 'Career Average'
            for cell in stats_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data
            for row_idx, (stat, value) in enumerate(career_avg.items(), 1):
                stats_table.rows[row_idx].cells[0].text = stat.replace('_', ' ').title()
                stats_table.rows[row_idx].cells[1].text = f"{value:.2f}"
        
        # AI Analysis
        doc.add_heading('AI-Powered Analysis', level=1)
        
        if ml_predictions and 'error' not in ml_predictions:
            doc.add_heading('Machine Learning Predictions', level=2)
            
            ml_para = doc.add_paragraph()
            ml_para.add_run(f"Current Impact Score: ").bold = True
            ml_para.add_run(f"{ml_predictions.get('current_impact_score', 0):.2f}\n")
            ml_para.add_run(f"Career Stage: ").bold = True
            ml_para.add_run(f"{ml_predictions.get('career_stage', 'Unknown')}\n")
            ml_para.add_run(f"Injury Risk: ").bold = True
            ml_para.add_run(f"{ml_predictions.get('injury_risk', 'Unknown')}")
        
        doc.add_heading('Expert Analysis', level=2)
        doc.add_paragraph(llm_analysis)
        
        doc.save(filename)
        logger.info(f"Generated player Word report: {filename}")
        
        return filename
    
    def _generate_team_pdf_report(self, team_data: Dict, ml_analysis: Dict, llm_analysis: str) -> str:
        """Generate PDF team analysis report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/team_report_{team_data['name'].replace(' ', '_')}_{timestamp}.pdf"
        
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        story.append(Paragraph(f"TEAM ANALYSIS REPORT: {team_data['name']}", title_style))
        story.append(Spacer(1, 20))
        
        # Team Overview
        story.append(Paragraph("Team Overview", styles['Heading2']))
        
        overview_data = [
            ['Team Name', team_data['name']],
            ['League', team_data.get('league', 'AFL')],
            ['Season', str(ml_analysis.get('season', 'Current'))],
        ]
        
        overview_table = Table(overview_data, colWidths=[2*inch, 3*inch])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(overview_table)
        story.append(Spacer(1, 20))
        
        # Performance Statistics
        if 'overall_stats' in ml_analysis:
            story.append(Paragraph("Performance Statistics", styles['Heading2']))
            
            stats = ml_analysis['overall_stats']
            stats_data = [['Metric', 'Value']]
            
            key_stats = [
                ('Total Matches', stats.get('total_matches', 0)),
                ('Wins', stats.get('wins', 0)),
                ('Losses', stats.get('losses', 0)),
                ('Win Percentage', f"{stats.get('win_percentage', 0):.1f}%"),
                ('Average Score For', f"{stats.get('avg_score_for', 0):.1f}"),
                ('Average Score Against', f"{stats.get('avg_score_against', 0):.1f}"),
                ('Average Margin', f"{stats.get('avg_margin', 0):.1f}"),
            ]
            
            for metric, value in key_stats:
                stats_data.append([metric, str(value)])
            
            stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(stats_table)
            story.append(Spacer(1, 20))
        
        # Expert Analysis
        story.append(Paragraph("Expert Analysis", styles['Heading2']))
        story.append(Paragraph(llm_analysis, styles['Normal']))
        
        doc.build(story)
        logger.info(f"Generated team PDF report: {filename}")
        
        return filename
    
    def _generate_team_excel_report(self, team_data: Dict, ml_analysis: Dict, llm_analysis: str) -> str:
        """Generate Excel team analysis report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/team_report_{team_data['name'].replace(' ', '_')}_{timestamp}.xlsx"
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Team Analysis"
        
        # Title
        ws.cell(row=1, column=1, value=f"Team Analysis Report: {team_data['name']}").font = Font(bold=True, size=16)
        
        # Team Overview
        ws.cell(row=3, column=1, value="Team Overview").font = Font(bold=True, size=14)
        
        overview_data = [
            ['Team Name', team_data['name']],
            ['League', team_data.get('league', 'AFL')],
            ['Season', str(ml_analysis.get('season', 'Current'))],
        ]
        
        for row, (attr, value) in enumerate(overview_data, 4):
            ws.cell(row=row, column=1, value=attr)
            ws.cell(row=row, column=2, value=value)
        
        # Performance Statistics
        if 'overall_stats' in ml_analysis:
            ws.cell(row=8, column=1, value="Performance Statistics").font = Font(bold=True, size=14)
            
            stats = ml_analysis['overall_stats']
            key_stats = [
                ('Total Matches', stats.get('total_matches', 0)),
                ('Wins', stats.get('wins', 0)),
                ('Losses', stats.get('losses', 0)),
                ('Win Percentage', f"{stats.get('win_percentage', 0):.1f}%"),
                ('Average Score For', f"{stats.get('avg_score_for', 0):.1f}"),
                ('Average Score Against', f"{stats.get('avg_score_against', 0):.1f}"),
                ('Average Margin', f"{stats.get('avg_margin', 0):.1f}"),
            ]
            
            for row, (metric, value) in enumerate(key_stats, 9):
                ws.cell(row=row, column=1, value=metric)
                ws.cell(row=row, column=2, value=value)
        
        # Expert Analysis
        ws.cell(row=18, column=1, value="Expert Analysis").font = Font(bold=True, size=14)
        ws.cell(row=19, column=1, value=llm_analysis)
        ws.cell(row=19, column=1).alignment = Alignment(wrap_text=True)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(filename)
        logger.info(f"Generated team Excel report: {filename}")
        
        return filename
    
    def _generate_team_word_report(self, team_data: Dict, ml_analysis: Dict, llm_analysis: str) -> str:
        """Generate Word team analysis report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/team_report_{team_data['name'].replace(' ', '_')}_{timestamp}.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Team Analysis Report: {team_data["name"]}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Team Overview
        doc.add_heading('Team Overview', level=1)
        
        overview_table = doc.add_table(rows=4, cols=2)
        overview_table.style = 'Table Grid'
        
        overview_data = [
            ['Attribute', 'Value'],
            ['Team Name', team_data['name']],
            ['League', team_data.get('league', 'AFL')],
            ['Season', str(ml_analysis.get('season', 'Current'))],
        ]
        
        for row_idx, (attr, value) in enumerate(overview_data):
            overview_table.rows[row_idx].cells[0].text = attr
            overview_table.rows[row_idx].cells[1].text = str(value)
            if row_idx == 0:  # Header row
                for cell in overview_table.rows[row_idx].cells:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # Performance Statistics
        if 'overall_stats' in ml_analysis:
            doc.add_heading('Performance Statistics', level=1)
            
            stats = ml_analysis['overall_stats']
            stats_table = doc.add_table(rows=8, cols=2)
            stats_table.style = 'Table Grid'
            
            stats_data = [
                ['Metric', 'Value'],
                ['Total Matches', str(stats.get('total_matches', 0))],
                ['Wins', str(stats.get('wins', 0))],
                ['Losses', str(stats.get('losses', 0))],
                ['Win Percentage', f"{stats.get('win_percentage', 0):.1f}%"],
                ['Average Score For', f"{stats.get('avg_score_for', 0):.1f}"],
                ['Average Score Against', f"{stats.get('avg_score_against', 0):.1f}"],
                ['Average Margin', f"{stats.get('avg_margin', 0):.1f}"],
            ]
            
            for row_idx, (metric, value) in enumerate(stats_data):
                stats_table.rows[row_idx].cells[0].text = metric
                stats_table.rows[row_idx].cells[1].text = value
                if row_idx == 0:  # Header row
                    for cell in stats_table.rows[row_idx].cells:
                        cell.paragraphs[0].runs[0].font.bold = True
        
        # Expert Analysis
        doc.add_heading('Expert Analysis', level=1)
        doc.add_paragraph(llm_analysis)
        
        doc.save(filename)
        logger.info(f"Generated team Word report: {filename}")
        
        return filename
    
    def _get_player_data(self, player_id: int) -> Dict:
        """Get comprehensive player data"""
        try:
            with db_manager.get_session() as session:
                # Get personal data
                personal_query = "SELECT * FROM players WHERE id = :player_id"
                personal_result = session.execute(personal_query, {'player_id': player_id}).fetchone()
                
                if not personal_result:
                    return None
                
                personal_data = dict(personal_result._mapping)
                
                # Get performance data
                performance_query = "SELECT * FROM player_performance WHERE player_id = :player_id ORDER BY year, games_played"
                performance_result = session.execute(performance_query, {'player_id': player_id}).fetchall()
                
                performance_data = [dict(row._mapping) for row in performance_result]
                
                return {
                    'personal': personal_data,
                    'performance': performance_data
                }
                
        except Exception as e:
            logger.error(f"Error getting player data: {e}")
            return None
    
    def _get_team_data(self, team_name: str, season: int = None) -> Dict:
        """Get comprehensive team data"""
        try:
            with db_manager.get_session() as session:
                # Get team basic info
                team_query = "SELECT * FROM teams WHERE name = :team_name LIMIT 1"
                team_result = session.execute(team_query, {'team_name': team_name}).fetchone()
                
                team_data = dict(team_result._mapping) if team_result else {'name': team_name}
                
                # Get recent matches
                season_filter = f"AND year = {season}" if season else "AND year >= 2020"
                matches_query = f"""
                SELECT * FROM matches 
                WHERE (team_1_name = :team_name OR team_2_name = :team_name)
                {season_filter}
                ORDER BY date DESC
                LIMIT 10
                """
                
                matches_result = session.execute(matches_query, {'team_name': team_name}).fetchall()
                recent_matches = [dict(row._mapping) for row in matches_result]
                
                team_data['recent_matches'] = recent_matches
                
                return team_data
                
        except Exception as e:
            logger.error(f"Error getting team data: {e}")
            return None
    
    def _get_league_data(self, season: int) -> Dict:
        """Get league-wide data for summary report"""
        try:
            with db_manager.get_session() as session:
                # Get all matches for the season
                matches_query = "SELECT * FROM matches WHERE year = :season ORDER BY date"
                matches_result = session.execute(matches_query, {'season': season}).fetchall()
                
                matches_data = [dict(row._mapping) for row in matches_result]
                
                # Get team standings
                teams_query = "SELECT DISTINCT name FROM teams WHERE league = 'AFLM' AND is_active = true"
                teams_result = session.execute(teams_query).fetchall()
                
                teams = [row[0] for row in teams_result]
                
                return {
                    'season': season,
                    'matches': matches_data,
                    'teams': teams
                }
                
        except Exception as e:
            logger.error(f"Error getting league data: {e}")
            return {}
    
    def _calculate_career_averages(self, perf_df: pd.DataFrame) -> Dict[str, float]:
        """Calculate career average statistics"""
        numeric_cols = ['kicks', 'marks', 'handballs', 'disposals', 'goals', 'behinds', 
                       'tackles', 'inside_50s', 'clearances', 'contested_possessions']
        
        averages = {}
        for col in numeric_cols:
            if col in perf_df.columns:
                averages[col] = perf_df[col].mean()
        
        return averages
    
    def _generate_league_pdf_report(self, league_data: Dict, season: int) -> str:
        """Generate league summary PDF report"""
        # Implementation for league report would go here
        # This is a placeholder for the full implementation
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/league_summary_{season}_{timestamp}.pdf"
        
        # Create basic report structure
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        title = Paragraph(f"AFL Season {season} Summary Report", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 30))
        
        content = Paragraph(f"League summary report for {season} season with {len(league_data.get('matches', []))} matches analyzed.", styles['Normal'])
        story.append(content)
        
        doc.build(story)
        logger.info(f"Generated league PDF report: {filename}")
        
        return filename
    
    def _generate_league_excel_report(self, league_data: Dict, season: int) -> str:
        """Generate league summary Excel report"""
        # Implementation for league Excel report would go here
        # This is a placeholder for the full implementation
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.report_dir}/league_summary_{season}_{timestamp}.xlsx"
        
        wb = Workbook()
        ws = wb.active
        ws.title = f"Season {season} Summary"
        
        ws.cell(row=1, column=1, value=f"AFL Season {season} Summary").font = Font(bold=True, size=16)
        ws.cell(row=3, column=1, value=f"Total Matches: {len(league_data.get('matches', []))}")
        ws.cell(row=4, column=1, value=f"Teams Analyzed: {len(league_data.get('teams', []))}")
        
        wb.save(filename)
        logger.info(f"Generated league Excel report: {filename}")
        
        return filename

# Global report generator instance
report_generator = ReportGenerator()